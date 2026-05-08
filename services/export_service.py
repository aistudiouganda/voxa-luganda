"""
Export service — generates transcript files in TXT, DOCX, PDF, SRT, VTT, JSON formats.
"""
import json
import logging
import io
from typing import Optional, List
from datetime import datetime

logger = logging.getLogger(__name__)


def format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def format_timestamp_srt(seconds: float) -> str:
    return format_timestamp(seconds)


def format_timestamp_vtt(seconds: float) -> str:
    return format_timestamp(seconds).replace(",", ".")


def generate_txt(segments: list, include_translation: bool = True) -> bytes:
    """Generate plain text transcript."""
    lines = []
    for seg in segments:
        time_str = f"[{format_secs(seg['start'])} → {format_secs(seg['end'])}]"
        speaker = seg.get("speakerLabel", f"Speaker {seg.get('speaker', 0) + 1}")
        lines.append(f"{time_str} {speaker}:")
        lines.append(seg["text"])
        if include_translation and seg.get("translation"):
            lines.append(f"  ({seg['translation']})")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def format_secs(s: float) -> str:
    m = int(s // 60)
    sec = int(s % 60)
    return f"{m:02d}:{sec:02d}"


def generate_srt(segments: list, use_translation: bool = False) -> bytes:
    """Generate SRT subtitle file."""
    lines = []
    for i, seg in enumerate(segments, 1):
        start = format_timestamp_srt(seg["start"])
        end = format_timestamp_srt(seg["end"])
        text = seg.get("translation", seg["text"]) if use_translation else seg["text"]
        speaker = seg.get("speakerLabel", "")
        lines.extend([
            str(i),
            f"{start} --> {end}",
            f"[{speaker}] {text}" if speaker else text,
            "",
        ])
    return "\n".join(lines).encode("utf-8")


def generate_vtt(segments: list, use_translation: bool = False) -> bytes:
    """Generate WebVTT subtitle file."""
    lines = ["WEBVTT", ""]
    for seg in segments:
        start = format_timestamp_vtt(seg["start"])
        end = format_timestamp_vtt(seg["end"])
        text = seg.get("translation", seg["text"]) if use_translation else seg["text"]
        speaker = seg.get("speakerLabel", "")
        lines.extend([
            f"{start} --> {end}",
            f"<v {speaker}>{text}" if speaker else text,
            "",
        ])
    return "\n".join(lines).encode("utf-8")


def generate_json(segments: list, metadata: dict = None) -> bytes:
    """Generate structured JSON output."""
    data = {
        "version": "1.0",
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "model": "voxa-luganda-v1",
        "metadata": metadata or {},
        "segments": segments,
    }
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def generate_docx(segments: list, title: str = "Transcript", include_translation: bool = True) -> bytes:
    """Generate formatted Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_run = title_para.runs[0]
        title_run.font.color.rgb = RGBColor(0x14, 0x14, 0x18)

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        meta.add_run(f" | Speakers: {len(set(s.get('speaker', 0) for s in segments))}")
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x6B, 0x7A)

        doc.add_paragraph()

        SPEAKER_COLORS = [
            RGBColor(0xE8, 0xA0, 0x20),  # gold
            RGBColor(0xC4, 0x54, 0x1B),  # terra
            RGBColor(0x1E, 0x6B, 0x45),  # savanna
            RGBColor(0x9B, 0x88, 0xFF),  # purple
        ]

        for seg in segments:
            speaker = seg.get("speakerLabel", f"Speaker {seg.get('speaker', 0) + 1}")
            speaker_idx = seg.get("speaker", 0)
            time_str = f"[{format_secs(seg['start'])} → {format_secs(seg['end'])}]"

            # Speaker label
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(f"{speaker}  {time_str}")
            speaker_run.bold = True
            speaker_run.font.size = Pt(9)
            speaker_run.font.color.rgb = SPEAKER_COLORS[speaker_idx % len(SPEAKER_COLORS)]
            speaker_para.paragraph_format.space_before = Pt(12)
            speaker_para.paragraph_format.space_after = Pt(2)

            # Luganda text
            text_para = doc.add_paragraph()
            text_para.add_run(seg["text"])
            text_para.paragraph_format.space_after = Pt(2)

            # Translation
            if include_translation and seg.get("translation"):
                trans_para = doc.add_paragraph()
                trans_run = trans_para.add_run(seg["translation"])
                trans_run.italic = True
                trans_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xB0)
                trans_run.font.size = Pt(10)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        logger.error("python-docx not installed")
        return generate_txt(segments, include_translation)


def generate_pdf(segments: list, title: str = "Transcript") -> bytes:
    """Generate PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=3*cm, rightMargin=3*cm)
        styles = getSampleStyleSheet()

        GOLD = colors.HexColor("#E8A020")
        TERRA = colors.HexColor("#C4541B")
        SAVANNA = colors.HexColor("#1E6B45")
        TEXT = colors.HexColor("#1C1C22")
        ASH = colors.HexColor("#6B6B7A")
        SPEAKER_COLORS_RPT = [GOLD, TERRA, SAVANNA, colors.HexColor("#9B88FF")]

        title_style = ParagraphStyle("Title", fontName="Helvetica-Bold", fontSize=20, textColor=TEXT, spaceAfter=6)
        meta_style = ParagraphStyle("Meta", fontName="Helvetica", fontSize=9, textColor=ASH, spaceAfter=20)
        speaker_style = ParagraphStyle("Speaker", fontName="Helvetica-Bold", fontSize=9, spaceAfter=3)
        text_style = ParagraphStyle("Text", fontName="Helvetica", fontSize=11, leading=16, textColor=TEXT, spaceAfter=3)
        trans_style = ParagraphStyle("Trans", fontName="Helvetica-Oblique", fontSize=10, textColor=ASH, leading=14, spaceAfter=10)

        story = [
            Paragraph(title, title_style),
            Paragraph(f"Generated by Voxa Luganda AI · {datetime.utcnow().strftime('%Y-%m-%d')}", meta_style),
        ]

        for seg in segments:
            speaker = seg.get("speakerLabel", "Speaker 1")
            speaker_idx = seg.get("speaker", 0)
            color = SPEAKER_COLORS_RPT[speaker_idx % len(SPEAKER_COLORS_RPT)]
            time_str = f"[{format_secs(seg['start'])} – {format_secs(seg['end'])}]"

            sp = ParagraphStyle(f"Spk{speaker_idx}", parent=speaker_style, textColor=color)
            story.append(Paragraph(f"{speaker}  <font color='grey'>{time_str}</font>", sp))
            story.append(Paragraph(seg["text"], text_style))
            if seg.get("translation"):
                story.append(Paragraph(seg["translation"], trans_style))

        doc.build(story)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        return generate_txt(segments)


def generate_full_text(segments: list) -> str:
    return " ".join(s.get("text", "") for s in segments)
